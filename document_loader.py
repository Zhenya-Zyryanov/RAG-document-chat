import csv
import io
from pathlib import Path


def load_document(file_path: str | Path) -> str:
    path = Path(file_path)
    suffix = path.suffix.lower()

    loaders = {
        ".txt": _load_txt,
        ".md":  _load_txt,
        ".pdf": _load_pdf,
        ".docx": _load_docx,
        ".csv":  _load_csv,
        ".xlsx": _load_xlsx,
        ".xls":  _load_xlsx,
    }

    loader = loaders.get(suffix)
    if loader is None:
        raise ValueError(
            f"Неподдерживаемый формат: '{suffix}'. "
            f"Поддерживаются: {', '.join(loaders)}"
        )

    text = loader(path)
    return text.strip()


def _load_txt(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="replace")


def _load_pdf(path: Path) -> str:
    try:
        import pdfplumber
    except ImportError:
        raise ImportError("Установи pdfplumber: pip install pdfplumber")

    pages = []
    with pdfplumber.open(str(path)) as pdf:
        for page in pdf.pages:
            parts = []

            # Обычный текст
            text = page.extract_text()
            if text:
                parts.append(text)

            for table in page.extract_tables():
                rows = []
                for row in table:
                    cells = [str(c).strip() if c is not None else "" for c in row]
                    if any(cells):
                        rows.append(" | ".join(cells))
                if rows:
                    parts.append("\n".join(rows))

            if parts:
                pages.append("\n\n".join(parts))

    return "\n\n".join(pages)

def _load_docx(path: Path) -> str:
    try:
        from docx import Document
    except ImportError:
        raise ImportError("python-docx: pip install python-docx")

    doc = Document(str(path))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

    # Также извлекаем текст из таблиц
    for table in doc.tables:
        for row in table.rows:
            row_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_texts:
                paragraphs.append(" | ".join(row_texts))

    return "\n".join(paragraphs)


def _load_csv(path: Path) -> str:
    lines = []
    with path.open(encoding="utf-8", errors="replace", newline="") as f:
        reader = csv.reader(f)
        for row in reader:
            lines.append(", ".join(row))
    return "\n".join(lines)


def _load_xlsx(path: Path) -> str:
    try:
        import openpyxl
    except ImportError:
        raise ImportError("openpyxl: pip install openpyxl")

    wb = openpyxl.load_workbook(str(path), read_only=True, data_only=True)
    sheets = []
    for sheet in wb.worksheets:
        rows = []
        for row in sheet.iter_rows(values_only=True):
            cells = [str(c) for c in row if c is not None]
            if cells:
                rows.append(" | ".join(cells))
        if rows:
            sheets.append(f"[Лист: {sheet.title}]\n" + "\n".join(rows))
    return "\n\n".join(sheets)
